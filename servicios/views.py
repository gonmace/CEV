from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db.models import Value, CharField, Count
from django.db.models.functions import Coalesce, Concat, NullIf, Trim, Lower
from django.db.models import Q
from django.core.files.base import ContentFile
from django.utils.text import slugify
from django.utils import timezone
from django.utils.safestring import mark_safe
from django.http import JsonResponse
from django.views.decorators.http import require_http_methods
import json
import re
import requests
import logging
from markdown import markdown
from .models import Servicio, ServicioImagen, CatalogoServicios
from .forms import ServicioForm

logger = logging.getLogger(__name__)

# OCR reader singleton — se inicializa una vez para no recargar el modelo en cada request
_ocr_reader = None

def _get_ocr_reader():
    global _ocr_reader
    if _ocr_reader is None:
        import easyocr, warnings
        with warnings.catch_warnings():
            warnings.filterwarnings('ignore', message='.*pin_memory.*')
            _ocr_reader = easyocr.Reader(['es', 'en'], gpu=False, verbose=False)
    return _ocr_reader

def extraer_texto_ocr(pdf_bytes_io):
    """Extrae texto y tablas de un PDF escaneado usando OCR local (easyocr + img2table)."""
    import fitz
    import tempfile, os
    from io import BytesIO

    pdf_bytes_io.seek(0)
    doc = fitz.open(stream=pdf_bytes_io.read(), filetype='pdf')
    reader = _get_ocr_reader()
    resultado = []

    for i, page in enumerate(list(doc)[:6]):
        mat = fitz.Matrix(1.5, 1.5)
        pix = page.get_pixmap(matrix=mat)
        img_bytes = pix.tobytes('png')

        # Texto libre
        try:
            lineas = reader.readtext(img_bytes, detail=0, paragraph=True)
            resultado.append(f"--- Página {i+1} ---")
            resultado.extend(lineas)
        except Exception as e:
            logger.warning(f"OCR texto página {i+1}: {e}")

        # Tablas estructuradas
        tmp_path = None
        try:
            from img2table.document import Image as Img2TableImage
            from img2table.ocr import EasyOCR as Img2EasyOCR
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                tmp.write(img_bytes)
                tmp_path = tmp.name
            img_doc = Img2TableImage(src=tmp_path)
            ocr_engine = Img2EasyOCR(lang=['es', 'en'])
            tablas = img_doc.extract_tables(ocr=ocr_engine, implicit_rows=True, borderless_tables=True)
            for tabla in tablas:
                resultado.append("[TABLA]")
                for fila in tabla.content.values():
                    resultado.append(" | ".join(c.value or '' for c in fila))
        except Exception as e:
            logger.warning(f"OCR tablas página {i+1}: {e}")
        finally:
            if tmp_path and os.path.exists(tmp_path):
                os.unlink(tmp_path)

    return '\n'.join(resultado).strip()[:8000]

N8N_WEBHOOK_SER_COHERENCIA_URL = 'https://n8n.magoreal.com/webhook/ser-coherencia'
N8N_WEBHOOK_SER_OBJETIVO_URL = 'https://n8n.magoreal.com/webhook/ser-objetivo'
N8N_WEBHOOK_SER_ALCANCE_URL = 'https://n8n.magoreal.com/webhook/ser-alcance'
N8N_WEBHOOK_SER_SECCIONES_URL = 'https://n8n.magoreal.com/webhook/ser-secciones'
N8N_WEBHOOK_SER_CLASIFICAR_URL = 'https://n8n.magoreal.com/webhook/ser-clasificar-alcance'
N8N_WEBHOOK_SER_EQUIPOS_EXTRACTOR_URL = 'https://n8n.magoreal.com/webhook/extractor-servicio'
N8N_WEBHOOK_SER_EQUIPOS_AJUSTAR_URL   = 'https://n8n.magoreal.com/webhook/ajustar-servicio'
N8N_WEBHOOK_SER_PDF_EXTRACTOR_URL = 'https://n8n.magoreal.com/webhook/pdf-vision'


def _categorias_json():
    """Devuelve el catálogo como JSON para cascada categoria→subcategoria en el template."""
    catalogo = CatalogoServicios.get_activo()
    datos = catalogo.datos if catalogo else []
    result = [
        {
            'id': cat['nombre'],
            'nombre': cat['nombre'],
            'subcategorias': [
                {
                    'id': sub['codigo'],
                    'codigo': sub['codigo'],
                    'nombre': sub['nombre'],
                    'definicion': sub.get('definicion', ''),
                    'alcance': sub.get('alcance', ''),
                    'descripcion': sub.get('descripcion', ''),
                }
                for sub in cat.get('subcategorias', [])
            ],
        }
        for cat in datos
    ]
    return json.dumps(result)


def _lookup_catalogo(subcategoria_codigo):
    """Devuelve (categoria_nombre, subcategoria_nombre) dado un código de subcategoría."""
    catalogo = CatalogoServicios.get_activo()
    if catalogo:
        for cat in catalogo.datos:
            for sub in cat.get('subcategorias', []):
                if sub['codigo'] == subcategoria_codigo:
                    return cat['nombre'], sub['nombre']
    return '', ''


@login_required
def lista_servicios_view(request):
    from django.core.paginator import Paginator

    sort_by = request.GET.get('sort_by', 'orden')
    order = request.GET.get('order', 'asc')

    valid_sort_fields = ['titulo', 'solicitante', 'fecha_creacion', 'publico', 'usuario', 'imagenes']
    if sort_by not in valid_sort_fields:
        sort_by = 'orden'
    if order not in ['asc', 'desc']:
        order = 'asc'

    per_page_options = [10, 20, 50, 100]
    per_page_raw = request.GET.get('per_page', str(per_page_options[0]))
    try:
        per_page = int(per_page_raw)
        if per_page not in per_page_options:
            per_page = per_page_options[0]
    except (TypeError, ValueError):
        per_page = per_page_options[0]

    base_qs = (
        Servicio.objects.filter(activo=True)
        .filter(Q(publico=True) | Q(creado_por=request.user))
        .select_related('creado_por')
        .annotate(
            usuario_full_name=Trim(Concat(
                Coalesce('creado_por__first_name', Value('', output_field=CharField())),
                Value(' ', output_field=CharField()),
                Coalesce('creado_por__last_name', Value('', output_field=CharField()))
            ))
        )
        .annotate(
            usuario_sort=Coalesce(
                NullIf('usuario_full_name', Value('', output_field=CharField())),
                'creado_por__username',
                Value('', output_field=CharField())
            )
        )
        .annotate(usuario_sort_lower=Lower('usuario_sort'))
        .annotate(num_imagenes=Count('imagenes', distinct=True))
    )

    # Separar completos (con contenido) de borradores (sin contenido)
    borradores_qs = base_qs.filter(contenido='').order_by('-fecha_actualizacion')
    servicios_qs = base_qs.exclude(contenido='')

    # Determinar paso siguiente para cada borrador
    _PASO_LABELS = {
        2: 'Objetivo',
        3: 'Alcance',
        4: 'Secciones',
        5: 'Consolidar',
    }
    borradores_ctx = []
    for b in borradores_qs:
        if b.secciones_generadas or b.secciones_editadas:
            paso = 5
        elif b.alcance_generado or b.alcance_editado:
            paso = 4
        elif b.objetivo:
            paso = 3
        else:
            paso = 2
        borradores_ctx.append({
            'obj': b,
            'resume_paso': paso,
            'resume_label': _PASO_LABELS[paso],
        })

    sort_field_map = {
        'titulo': 'titulo',
        'solicitante': 'solicitante',
        'fecha_creacion': 'fecha_creacion',
        'publico': 'publico',
        'usuario': 'usuario_sort_lower',
        'imagenes': 'num_imagenes',
    }
    order_field = sort_field_map.get(sort_by, 'orden')
    order_by = f'-{order_field}' if order == 'desc' else order_field
    servicios_qs = servicios_qs.order_by(order_by)

    paginator = Paginator(servicios_qs, per_page)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    query_params = request.GET.copy()
    if 'page' in query_params:
        query_params.pop('page')
    base_query = query_params.urlencode()

    from django.contrib.auth.models import User as AuthUser
    total_usuarios = AuthUser.objects.filter(is_active=True).count()

    catalogo = CatalogoServicios.get_activo()
    datos_catalogo = catalogo.datos if catalogo else []
    total_categorias = len(datos_catalogo)
    total_subcategorias = sum(len(c.get('subcategorias', [])) for c in datos_catalogo)

    return render(request, 'servicios/index.html', {
        'servicios': page_obj.object_list,
        'page_obj': page_obj,
        'base_query': base_query,
        'sort_by': sort_by,
        'order': order,
        'per_page': per_page,
        'per_page_options': per_page_options,
        'total_usuarios': total_usuarios,
        'total_categorias': total_categorias,
        'total_subcategorias': total_subcategorias,
        'datos_catalogo': datos_catalogo,
        'borradores': borradores_ctx,
    })


@login_required
def crear_servicio_view(request):
    if request.method == 'POST':
        form = ServicioForm(request.POST)
        if form.is_valid():
            codigo = form.cleaned_data['subcategoria_codigo']
            cat_nombre, sub_nombre = _lookup_catalogo(codigo)

            catalogo = CatalogoServicios.get_activo()
            subcategorias_categoria = []
            if catalogo:
                for cat in catalogo.datos:
                    if cat['nombre'] == cat_nombre:
                        subcategorias_categoria = [
                            {
                                'codigo': sub['codigo'],
                                'nombre': sub['nombre'],
                                'descripcion': sub.get('descripcion', ''),
                                'intencion': sub.get('intencion', ''),
                            }
                            for sub in cat.get('subcategorias', [])
                        ]
                        break
            payload = {
                'titulo': form.cleaned_data['titulo'],
                'descripcion': form.cleaned_data['descripcion'],
                'categoria_nombre': cat_nombre,
                'subcategoria_codigo': codigo,
                'subcategoria_nombre': sub_nombre,
                'subcategorias_disponibles': subcategorias_categoria,
            }
            sugerencia = None
            bypass = request.POST.get('bypass_coherencia')
            if not bypass:
                try:
                    resp = requests.post(
                        N8N_WEBHOOK_SER_COHERENCIA_URL,
                        json=payload,
                        timeout=30,
                    )
                    resp.raise_for_status()
                    data = resp.json()
                    coherente = data.get('coherente', True)
                    asignacion = data.get('asignacion', True)
                    if not coherente or not asignacion:
                        sugerencia = {
                            'coherente': coherente,
                            'asignacion': asignacion,
                            'razon_coherencia': data.get('razon_coherencia', ''),
                            'razon_asignacion': data.get('razon_asignacion', ''),
                            'subcategoria_sugerida': data.get('subcategoria_sugerida', ''),
                        }
                        return render(request, 'servicios/crear_servicio.html', {
                            'form': form,
                            'categorias_json': _categorias_json(),
                            'sugerencia': sugerencia,
                            'bypass_coherencia': True,
                        })
                except requests.exceptions.RequestException as e:
                    logger.warning(f"No se pudo contactar el webhook de coherencia: {e}")

            servicio = form.save(commit=False)
            servicio.categoria_nombre = cat_nombre
            servicio.subcategoria_nombre = sub_nombre
            servicio.creado_por = request.user
            servicio.save()
            return redirect('servicios:paso2_objetivo', servicio_id=servicio.id)
    else:
        form = ServicioForm()
    return render(request, 'servicios/crear_servicio.html', {
        'form': form,
        'categorias_json': _categorias_json(),
    })


@login_required
def paso2_objetivo_view(request, servicio_id):
    servicio = get_object_or_404(Servicio, id=servicio_id, activo=True, creado_por=request.user)

    if request.method == 'POST':
        objetivo = request.POST.get('objetivo', '').strip()
        servicio.objetivo = objetivo
        servicio.save(update_fields=['objetivo'])
        return redirect('servicios:paso3_alcance', servicio_id=servicio.id)

    return render(request, 'servicios/paso2_objetivo.html', {'servicio': servicio})


@login_required
@require_http_methods(['POST'])
def generar_objetivo_ajax(request, servicio_id):
    servicio = get_object_or_404(Servicio, id=servicio_id, activo=True, creado_por=request.user)
    try:
        body = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        body = {}

    payload = {
        'titulo': servicio.titulo,
        'descripcion': servicio.descripcion,
        'categoria_nombre': servicio.categoria_nombre,
        'subcategoria_nombre': servicio.subcategoria_nombre,
        'historial': body.get('historial', []),
    }
    try:
        resp = requests.post(
            N8N_WEBHOOK_SER_OBJETIVO_URL,
            json=payload,
            timeout=60,
        )
        resp.raise_for_status()
        return JsonResponse(resp.json())
    except requests.exceptions.RequestException as e:
        logger.warning(f"Error webhook objetivo: {e}")
        return JsonResponse({'error': 'No se pudo conectar con el generador de objetivos.'}, status=503)


@login_required
def paso3_alcance_view(request, servicio_id):
    servicio = get_object_or_404(Servicio, id=servicio_id, activo=True, creado_por=request.user)

    if request.method == 'POST':
        alcance_raw = request.POST.get('alcance', '').strip().replace('\x00', '')
        tiene_header = '## alcance' in alcance_raw.lower()
        tiene_tabla = '|' in alcance_raw
        if alcance_raw and (not tiene_header or not tiene_tabla):
            messages.warning(request, 'El alcance no tiene el formato esperado. Puedes editarlo y volver a guardar.')
        if not servicio.alcance_generado:
            servicio.alcance_generado = alcance_raw
        servicio.alcance_editado = alcance_raw
        servicio.save(update_fields=['alcance_generado', 'alcance_editado'])
        return redirect('servicios:paso4_secciones', servicio_id=servicio.id)

    return render(request, 'servicios/paso3_alcance.html', {'servicio': servicio})


@login_required
def paso4_secciones_view(request, servicio_id):
    servicio = get_object_or_404(Servicio, id=servicio_id, activo=True, creado_por=request.user)

    if request.method == 'POST':
        secciones_raw = request.POST.get('secciones', '').strip().replace('\x00', '')
        if secciones_raw:
            if not servicio.secciones_generadas:
                servicio.secciones_generadas = secciones_raw
            servicio.secciones_editadas = secciones_raw
            servicio.save(update_fields=['secciones_generadas', 'secciones_editadas'])
        return redirect('servicios:paso5_consolidar', servicio_id=servicio.id)

    return render(request, 'servicios/paso4_secciones.html', {'servicio': servicio})


@login_required
def paso5_consolidar_view(request, servicio_id):
    servicio = get_object_or_404(Servicio, id=servicio_id, activo=True, creado_por=request.user)

    if request.method == 'POST':
        contenido_raw = request.POST.get('contenido', '').strip().replace('\x00', '')
        servicio.contenido = contenido_raw
        servicio.save(update_fields=['contenido'])
        messages.success(request, f'Servicio "{servicio.titulo}" creado exitosamente.')
        return redirect('servicios:ver_servicio', servicio_id=servicio.id)

    # Construir contenido consolidado
    partes = []
    if servicio.objetivo:
        partes.append(f"## Objetivo\n\n{servicio.objetivo}")
    alcance = servicio.alcance_editado or servicio.alcance_generado
    if alcance:
        # Los ## internos del alcance (excepto el título principal) se bajan a ###
        lineas_alc, primera_h2 = [], True
        for l in alcance.splitlines():
            if l.startswith('## '):
                if primera_h2:
                    primera_h2 = False
                    lineas_alc.append(l)
                else:
                    lineas_alc.append('### ' + l[3:])
            else:
                lineas_alc.append(l)
        partes.append('\n'.join(lineas_alc))
    secciones = servicio.secciones_editadas or servicio.secciones_generadas
    if secciones:
        partes.append(secciones)
    contenido_md = servicio.contenido or '\n\n'.join(partes)

    # Aplicar Title Case a encabezados ##
    _conjunciones = {'y', 'e', 'o', 'u', 'a', 'de', 'del', 'el', 'la', 'los', 'las',
                     'un', 'una', 'en', 'con', 'por', 'para', 'que', 'si', 'al', 'su', 'sus'}

    def _title_case(texto):
        palabras = texto.split()
        return ' '.join(
            p[0].upper() + p[1:].lower() if (i == 0 or p.lower() not in _conjunciones) else p.lower()
            for i, p in enumerate(palabras)
        )

    lineas = []
    for linea in contenido_md.splitlines():
        if linea.startswith('## '):
            lineas.append('## ' + _title_case(linea[3:]))
        else:
            lineas.append(linea)
    contenido_md = '\n'.join(lineas)

    preview_html = mark_safe(markdown(contenido_md, extensions=['extra']))

    return render(request, 'servicios/paso5_consolidar.html', {
        'servicio': servicio,
        'contenido_md': contenido_md,
        'preview_html': preview_html,
    })


@login_required
@require_http_methods(['POST'])
def clasificar_alcance_ajax(request, servicio_id):
    servicio = get_object_or_404(Servicio, id=servicio_id, activo=True, creado_por=request.user)

    catalogo = CatalogoServicios.get_activo()
    intencion_raw = ''
    if catalogo:
        for cat in catalogo.datos:
            for sub in cat.get('subcategorias', []):
                if sub['codigo'] == servicio.subcategoria_codigo:
                    intencion_raw = sub.get('intencion', '')
                    break
    intenciones = [i.strip() for i in intencion_raw.split(',') if i.strip()]

    payload = {
        'titulo': servicio.titulo,
        'descripcion': servicio.descripcion,
        'objetivo': servicio.objetivo,
        'subcategoria_nombre': servicio.subcategoria_nombre,
        'intenciones': intenciones,
    }
    try:
        resp = requests.post(
            N8N_WEBHOOK_SER_CLASIFICAR_URL,
            json=payload,
            timeout=30,
        )
        resp.raise_for_status()
        data = resp.json()
        if isinstance(data, list):
            data = data[0] if data else {}
        result = data.get('output', data) if isinstance(data, dict) else {}
        return JsonResponse(result)
    except Exception as e:
        logger.warning(f"Error clasificador: {e}")
        return JsonResponse({'error': 'No se pudo clasificar el servicio.'}, status=503)


def _parsear_paginas(paginas_str, total_paginas):
    """Convierte "1,3,5" o "1-3" o "2,4-5" a lista de índices 0-based, máximo 3."""
    indices = set()
    for parte in paginas_str.replace(' ', '').split(','):
        if '-' in parte:
            extremos = parte.split('-', 1)
            try:
                ini, fin = int(extremos[0]), int(extremos[1])
                indices.update(range(ini, fin + 1))
            except ValueError:
                pass
        else:
            try:
                indices.add(int(parte))
            except ValueError:
                pass
    # Convertir a 0-based y validar contra total
    validos = sorted(i - 1 for i in indices if 1 <= i <= total_paginas)
    if not validos:
        validos = list(range(min(3, total_paginas)))
    return validos[:3]


def _vision_paginas(pdf_bytes_io, indices_paginas, nombre, servicio):
    """Ejecuta Vision sobre las páginas indicadas (índices 0-based)."""
    import fitz, base64
    pdf_bytes_io.seek(0)
    doc = fitz.open(stream=pdf_bytes_io.read(), filetype='pdf')
    imagenes = []
    for i in indices_paginas:
        if i < len(doc):
            mat = fitz.Matrix(1.2, 1.2)
            pix = doc[i].get_pixmap(matrix=mat)
            imagenes.append(base64.b64encode(pix.tobytes('jpeg', jpg_quality=70)).decode())
    payload_vision = {
        'nombre': nombre,
        'imagenes': imagenes,
        'titulo_servicio': servicio.titulo,
        'descripcion_servicio': servicio.descripcion or '',
        'subcategoria': servicio.subcategoria_nombre or '',
    }
    resp_v = requests.post(
        N8N_WEBHOOK_SER_PDF_EXTRACTOR_URL,
        json=payload_vision,
        timeout=90,
    )
    resp_v.raise_for_status()
    data_v = resp_v.json()
    if isinstance(data_v, list):
        data_v = data_v[0] if data_v else {}
    return data_v.get('output', data_v) if isinstance(data_v, dict) else {}


@login_required
@require_http_methods(['POST'])
def extraer_equipo_ajax(request, servicio_id):
    servicio = get_object_or_404(Servicio, id=servicio_id, activo=True, creado_por=request.user)

    tipo = request.POST.get('tipo', '').strip()
    nombre = request.POST.get('nombre', '').strip()
    contenido = request.POST.get('contenido', '').strip()
    paginas_str = request.POST.get('paginas', '').strip()

    texto = ''
    if tipo == 'url':
        if not contenido:
            return JsonResponse({'error': 'URL vacía.'}, status=400)
        try:
            r = requests.get(contenido, timeout=15, headers={'User-Agent': 'Mozilla/5.0'})
            content_type = r.headers.get('Content-Type', '').lower()
            is_pdf = 'pdf' in content_type or contenido.lower().split('?')[0].endswith('.pdf')

            if is_pdf:
                # Tratar como PDF
                from io import BytesIO
                import pdfplumber
                pdf_bytes = r.content
                if len(pdf_bytes) > 5 * 1024 * 1024:
                    return JsonResponse({'error': 'El PDF de la URL no debe superar 5 MB.'}, status=400)
                pdf_bytes_io = BytesIO(pdf_bytes)
                total_paginas = 0
                try:
                    with pdfplumber.open(pdf_bytes_io) as pdf:
                        total_paginas = len(pdf.pages)
                        if total_paginas > 3 and not paginas_str:
                            return JsonResponse({'necesita_paginas': True, 'total_paginas': total_paginas})
                        indices = _parsear_paginas(paginas_str, total_paginas) if paginas_str else list(range(min(3, total_paginas)))
                        texto = '\n'.join(pdf.pages[i].extract_text() or '' for i in indices).strip()[:6000]
                except Exception:
                    texto = ''

                if not texto:
                    # PDF escaneado → Vision (gpt-4o-mini)
                    try:
                        import fitz
                        if total_paginas == 0:
                            pdf_bytes_io.seek(0)
                            doc_count = fitz.open(stream=pdf_bytes_io.read(), filetype='pdf')
                            total_paginas = len(doc_count)
                        if not paginas_str and total_paginas > 3:
                            return JsonResponse({'necesita_paginas': True, 'total_paginas': total_paginas})
                        indices = _parsear_paginas(paginas_str, total_paginas) if paginas_str else list(range(min(3, total_paginas)))
                        result_v = _vision_paginas(pdf_bytes_io, indices, nombre, servicio)
                        return JsonResponse(result_v)
                    except Exception as e:
                        logger.exception(f"Error Vision PDF desde URL: {e}")
                        return JsonResponse({'error': 'PDF escaneado: no se pudo procesar.'}, status=422)
            else:
                html = r.text
                html = re.sub(r'<script[\s\S]*?>[\s\S]*?</script>', '', html, flags=re.IGNORECASE)
                html = re.sub(r'<style[\s\S]*?>[\s\S]*?</style>', '', html, flags=re.IGNORECASE)
                html = re.sub(r'<[^>]+>', '', html)
                html = re.sub(r'\s{2,}', ' ', html).strip()
                texto = html[:6000]
        except Exception as e:
            return JsonResponse({'error': f'No se pudo obtener la URL: {e}'}, status=400)

    elif tipo == 'pdf':
        pdf_file = request.FILES.get('archivo')
        if not pdf_file:
            return JsonResponse({'error': 'No se proporcionó archivo PDF.'}, status=400)
        if pdf_file.size > 5 * 1024 * 1024:
            return JsonResponse({'error': 'El PDF no debe superar 5 MB.'}, status=400)
        total_paginas = 0
        try:
            import pdfplumber
            from io import BytesIO
            pdf_bytes_io = BytesIO(pdf_file.read())
            with pdfplumber.open(pdf_bytes_io) as pdf:
                total_paginas = len(pdf.pages)
                if total_paginas > 3 and not paginas_str:
                    return JsonResponse({'necesita_paginas': True, 'total_paginas': total_paginas})
                indices = _parsear_paginas(paginas_str, total_paginas) if paginas_str else list(range(min(3, total_paginas)))
                texto = '\n'.join(pdf.pages[i].extract_text() or '' for i in indices)
            texto = texto.strip()[:6000]
        except Exception as e:
            logger.exception(f"Error extrayendo PDF: {e}")
            return JsonResponse({'error': f'Error extrayendo PDF: {e}'}, status=500)

        # PDF escaneado → Vision (gpt-4o-mini)
        if not texto:
            try:
                import fitz
                if total_paginas == 0:
                    pdf_bytes_io.seek(0)
                    doc_count = fitz.open(stream=pdf_bytes_io.read(), filetype='pdf')
                    total_paginas = len(doc_count)
                if not paginas_str and total_paginas > 3:
                    return JsonResponse({'necesita_paginas': True, 'total_paginas': total_paginas})
                indices = _parsear_paginas(paginas_str, total_paginas) if paginas_str else list(range(min(3, total_paginas)))
                result_v = _vision_paginas(pdf_bytes_io, indices, nombre, servicio)
                return JsonResponse(result_v)
            except Exception as e:
                logger.exception(f"Error Vision PDF: {e}")
                return JsonResponse({'error': 'PDF escaneado: no se pudo procesar.'}, status=422)

    elif tipo == 'texto':
        if not contenido:
            return JsonResponse({'error': 'El texto está vacío.'}, status=400)
        texto = contenido[:6000]

    else:
        return JsonResponse({'error': 'Tipo inválido.'}, status=400)

    if not texto:
        return JsonResponse({'error': 'No se pudo extraer texto. Verifica que el archivo o URL tenga contenido.'}, status=400)

    payload = {
        'nombre': nombre,
        'tipo_fuente': tipo,
        'contenido': texto,
        'titulo_servicio': servicio.titulo,
        'descripcion_servicio': servicio.descripcion or '',
        'subcategoria': servicio.subcategoria_nombre or '',
    }
    try:
        resp = requests.post(
            N8N_WEBHOOK_SER_EQUIPOS_EXTRACTOR_URL,
            json=payload,
            timeout=60,
        )
        resp.raise_for_status()
        data = resp.json()
        if isinstance(data, list):
            data = data[0] if data else {}
        result = data.get('output', data) if isinstance(data, dict) else {}
        return JsonResponse(result)
    except Exception as e:
        logger.exception(f"Error extractor equipos: {e}")
        return JsonResponse({'error': 'No se pudo conectar con el extractor de equipos.'}, status=503)


@login_required
@require_http_methods(['POST'])
def generar_alcance_ajax(request, servicio_id):
    servicio = get_object_or_404(Servicio, id=servicio_id, activo=True, creado_por=request.user)
    try:
        body = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        body = {}

    # Obtener intencion de la subcategoría en el catálogo
    catalogo = CatalogoServicios.get_activo()
    intencion_raw = ''
    if catalogo:
        for cat in catalogo.datos:
            for sub in cat.get('subcategorias', []):
                if sub['codigo'] == servicio.subcategoria_codigo:
                    intencion_raw = sub.get('intencion', '')
                    break
    intenciones = [i.strip() for i in intencion_raw.split(',') if i.strip()]

    # Construir contexto_web desde múltiples referencias
    referencias = body.get('referencias', [])
    if not referencias:
        # Compatibilidad con formato anterior
        url_ref = body.get('url_referencia', '')
        texto_ref = body.get('texto_referencia', '')
        if url_ref or texto_ref:
            referencias = [{'nombre': '', 'url_referencia': url_ref, 'texto_referencia': texto_ref}]

    partes_contexto = []
    for ref in referencias:
        nombre = ref.get('nombre', '').strip()
        url = ref.get('url_referencia', '').strip()
        texto = ref.get('texto_referencia', '').strip()
        if url:
            try:
                r = requests.get(url, timeout=15, headers={'User-Agent': 'Mozilla/5.0'})
                html = r.text
                html = re.sub(r'<script[\s\S]*?>[\s\S]*?</script>', '', html, flags=re.IGNORECASE)
                html = re.sub(r'<style[\s\S]*?>[\s\S]*?</style>', '', html, flags=re.IGNORECASE)
                html = re.sub(r'<[^>]+>', '', html)
                html = re.sub(r'\s{2,}', ' ', html).strip()
                header = f"[{nombre}]:" if nombre else f"Referencia (URL: {url}):"
                partes_contexto.append(f"{header}\n{html[:4000]}")
            except Exception as fetch_err:
                logger.warning(f"No se pudo obtener URL {url}: {fetch_err}")
        elif texto:
            header = f"[{nombre}]:" if nombre else "Referencia:"
            partes_contexto.append(f"{header}\n{texto[:4000]}")
    contexto_web = '\n\n---\n\n'.join(partes_contexto)

    # Usar estructura provista por el frontend (ya fue clasificado en paso previo)
    # Solo llamar al clasificador como fallback si no viene en el body
    estructura = body.get('estructura', '').strip()
    if not estructura:
        try:
            payload_clasificar = {
                'titulo': servicio.titulo,
                'descripcion': servicio.descripcion,
                'objetivo': servicio.objetivo,
                'subcategoria_nombre': servicio.subcategoria_nombre,
                'intenciones': intenciones,
            }
            resp_clas = requests.post(
                N8N_WEBHOOK_SER_CLASIFICAR_URL,
                json=payload_clasificar,
                timeout=30,
            )
            if resp_clas.ok and resp_clas.text.strip():
                data_clas = resp_clas.json()
                if isinstance(data_clas, list):
                    data_clas = data_clas[0] if data_clas else {}
                result_clas = data_clas.get('output', data_clas) if isinstance(data_clas, dict) else {}
                estructura = result_clas.get('estructura', '')
        except Exception as e:
            logger.warning(f"Clasificador alcance falló (se continúa sin estructura): {e}")

    equipos = body.get('equipos', [])

    # Persistir equipos en el modelo si vienen en el payload
    if equipos:
        servicio.equipos = equipos
        servicio.save(update_fields=['equipos'])

    payload = {
        'titulo': servicio.titulo,
        'descripcion': servicio.descripcion,
        'objetivo': servicio.objetivo,
        'categoria_nombre': servicio.categoria_nombre,
        'subcategoria_nombre': servicio.subcategoria_nombre,
        'intenciones': intenciones,
        'estructura': estructura,
        'modalidad': body.get('modalidad', 'EVENTUAL'),
        'equipos': equipos,
        'historial': body.get('historial', []),
        'contexto_web': contexto_web,
    }
    try:
        resp = requests.post(
            N8N_WEBHOOK_SER_ALCANCE_URL,
            json=payload,
            timeout=120,
        )
        resp.raise_for_status()
        if not resp.text.strip():
            logger.error(f"n8n devolvió respuesta vacía (status {resp.status_code})")
            return JsonResponse({'error': 'El generador no respondió. Intenta de nuevo.'}, status=503)
        data = resp.json()
        # n8n puede devolver lista o dict
        if isinstance(data, list):
            data = data[0] if data else {}
        # desempaquetar campo 'output' si existe
        result = data.get('output', data) if isinstance(data, dict) else {}
        # Guardar alcance inmediatamente en DB cuando n8n lo genera
        if result.get('tipo') == 'alcance' and result.get('alcance'):
            if not servicio.alcance_generado:
                servicio.alcance_generado = result['alcance'].replace('\x00', '')
                servicio.save(update_fields=['alcance_generado'])
        return JsonResponse(result)
    except Exception as e:
        logger.exception(f"Error webhook alcance: {e}")
        return JsonResponse({'error': 'No se pudo conectar con el generador de alcance.'}, status=503)


@login_required
@require_http_methods(['POST'])
def generar_secciones_ajax(request, servicio_id):
    servicio = get_object_or_404(Servicio, id=servicio_id, activo=True, creado_por=request.user)
    try:
        body = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        body = {}

    estructura = body.get('estructura', '').strip()
    if not estructura:
        catalogo = CatalogoServicios.get_activo()
        intencion_raw = ''
        if catalogo:
            for cat in catalogo.datos:
                for sub in cat.get('subcategorias', []):
                    if sub['codigo'] == servicio.subcategoria_codigo:
                        intencion_raw = sub.get('intencion', '')
                        break
        intenciones = [i.strip() for i in intencion_raw.split(',') if i.strip()]
        try:
            resp_clas = requests.post(
                N8N_WEBHOOK_SER_CLASIFICAR_URL,
                json={
                    'titulo': servicio.titulo,
                    'descripcion': servicio.descripcion,
                    'objetivo': servicio.objetivo,
                    'subcategoria_nombre': servicio.subcategoria_nombre,
                    'intenciones': intenciones,
                },
                timeout=30,
            )
            if resp_clas.ok and resp_clas.text.strip():
                data_clas = resp_clas.json()
                if isinstance(data_clas, list):
                    data_clas = data_clas[0] if data_clas else {}
                result_clas = data_clas.get('output', data_clas) if isinstance(data_clas, dict) else {}
                estructura = result_clas.get('estructura', '')
        except Exception as e:
            logger.warning(f"Clasificador secciones falló: {e}")

    payload = {
        'titulo': servicio.titulo,
        'descripcion': servicio.descripcion,
        'objetivo': servicio.objetivo,
        'categoria_nombre': servicio.categoria_nombre,
        'subcategoria_nombre': servicio.subcategoria_nombre,
        'estructura': estructura,
        'modalidad': body.get('modalidad', 'EVENTUAL'),
        'equipos': body.get('equipos', servicio.equipos or []),
        'alcance_generado': body.get('alcance_generado', servicio.alcance_generado or ''),
        'historial': body.get('historial', []),
        'contexto_web': body.get('contexto_web', ''),
    }
    try:
        resp = requests.post(
            N8N_WEBHOOK_SER_SECCIONES_URL,
            json=payload,
            timeout=120,
        )
        resp.raise_for_status()
        if not resp.text.strip():
            return JsonResponse({'error': 'El generador no respondió. Intenta de nuevo.'}, status=503)
        data = resp.json()
        if isinstance(data, list):
            data = data[0] if data else {}
        result = data.get('output', data) if isinstance(data, dict) else {}
        # Guardar secciones inmediatamente en DB cuando n8n las genera
        if result.get('tipo') == 'secciones' and result.get('secciones'):
            if not servicio.secciones_generadas:
                servicio.secciones_generadas = result['secciones'].replace('\x00', '')
                servicio.save(update_fields=['secciones_generadas'])
        return JsonResponse(result)
    except Exception as e:
        logger.exception(f"Error webhook secciones: {e}")
        return JsonResponse({'error': 'No se pudo conectar con el generador de secciones.'}, status=503)


@login_required
def ver_servicio_view(request, servicio_id):
    servicio = get_object_or_404(
        Servicio.objects.select_related('creado_por').prefetch_related('imagenes'),
        id=servicio_id, activo=True
    )
    if not (servicio.publico or servicio.creado_por == request.user):
        messages.error(request, 'No tienes permisos para ver este servicio.')
        return redirect('servicios:lista_servicios')

    es_propietario = servicio.creado_por == request.user
    contenido_md = servicio.contenido or ''
    preview_html = mark_safe(markdown(contenido_md, extensions=['extra']))
    tiene_cantidad = bool(servicio.cantidad and servicio.cantidad.strip())

    return render(request, 'servicios/ver_servicio.html', {
        'servicio': servicio,
        'es_propietario': es_propietario,
        'preview_html': preview_html,
        'contenido_md_json': json.dumps(contenido_md),
        'tiene_cantidad': tiene_cantidad,
    })


@login_required
def editar_servicio_view(request, servicio_id):
    servicio = get_object_or_404(Servicio, id=servicio_id, activo=True)

    if servicio.creado_por != request.user:
        messages.error(request, 'Solo puedes editar servicios de tu cuenta.')
        return redirect('servicios:ver_servicio', servicio_id=servicio_id)

    if request.method == 'POST':
        form = ServicioForm(request.POST, instance=servicio)
        if form.is_valid():
            codigo = form.cleaned_data['subcategoria_codigo']
            cat_nombre, sub_nombre = _lookup_catalogo(codigo)
            servicio = form.save(commit=False)
            servicio.categoria_nombre = cat_nombre
            servicio.subcategoria_nombre = sub_nombre
            servicio.save()
            messages.success(request, 'Servicio actualizado correctamente.')
            return redirect('servicios:ver_servicio', servicio_id=servicio.id)
    else:
        form = ServicioForm(instance=servicio)

    return render(request, 'servicios/editar_servicio.html', {
        'form': form,
        'servicio': servicio,
        'categorias_json': _categorias_json(),
    })


@login_required
@require_http_methods(['POST'])
def guardar_contenido_ajax(request, servicio_id):
    servicio = get_object_or_404(Servicio, id=servicio_id, activo=True)
    if servicio.creado_por != request.user:
        return JsonResponse({'error': 'Sin permiso'}, status=403)
    body = json.loads(request.body)
    contenido = body.get('contenido', '').strip().replace('\x00', '')
    servicio.contenido = contenido
    servicio.save(update_fields=['contenido'])
    return JsonResponse({'ok': True})


@login_required
def eliminar_servicio_view(request, servicio_id):
    servicio = get_object_or_404(Servicio, id=servicio_id, activo=True)

    if servicio.creado_por != request.user:
        is_ajax = (
            request.headers.get('X-Requested-With') == 'XMLHttpRequest'
            or request.content_type == 'application/json'
        )
        if is_ajax:
            return JsonResponse({'error': 'Solo puedes eliminar servicios de tu cuenta.'}, status=403)
        messages.error(request, 'Solo puedes eliminar servicios de tu cuenta.')
        return redirect('servicios:lista_servicios')

    if request.method == 'POST':
        is_ajax = (
            request.headers.get('X-Requested-With') == 'XMLHttpRequest'
            or request.content_type == 'application/json'
        )
        if servicio.archivo:
            servicio.archivo.delete(save=False)
        servicio.activo = False
        servicio.fecha_eliminacion = timezone.now()
        servicio.save()
        if is_ajax:
            return JsonResponse({'success': True, 'message': f'Servicio "{servicio.titulo}" eliminado correctamente.'})
        messages.success(request, f'Servicio "{servicio.titulo}" eliminado correctamente.')
        return redirect('servicios:lista_servicios')

    return render(request, 'servicios/eliminar_servicio.html', {'servicio': servicio})


@login_required
@require_http_methods(["POST"])
def toggle_publico_view(request, servicio_id):
    servicio = get_object_or_404(Servicio, id=servicio_id, activo=True)
    if servicio.creado_por != request.user:
        return JsonResponse({'error': 'Sin permisos.'}, status=403)
    servicio.publico = not servicio.publico
    servicio.save(update_fields=['publico'])
    return JsonResponse({'publico': servicio.publico})


@login_required
def exportar_servicio_word_view(request, servicio_id):
    from django.http import HttpResponse
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from bs4 import BeautifulSoup
    import io

    servicio = get_object_or_404(Servicio, id=servicio_id, activo=True)
    if not (servicio.publico or servicio.creado_por == request.user):
        messages.error(request, 'Sin permisos para exportar este servicio.')
        return redirect('servicios:ver_servicio', servicio_id=servicio_id)

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # — Cabecera —
    if servicio.subcategoria_codigo:
        p = doc.add_paragraph()
        run = p.add_run(servicio.subcategoria_codigo)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1d, 0x4e, 0xd8)
        run.bold = True

    h = doc.add_heading(servicio.titulo, level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)

    meta_parts = []
    if servicio.categoria_nombre:
        meta_parts.append(servicio.categoria_nombre)
    if servicio.subcategoria_nombre:
        meta_parts.append(servicio.subcategoria_nombre)
    if servicio.solicitante:
        meta_parts.append(servicio.solicitante)
    if meta_parts:
        p = doc.add_paragraph(' › '.join(meta_parts))
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    doc.add_paragraph()  # separador

    # — Contenido —
    if servicio.contenido and servicio.contenido.strip():
        contenido_html = markdown(servicio.contenido, extensions=['extra'])
        soup = BeautifulSoup(contenido_html, 'html.parser')

        def add_run_formatted(para, node):
            if not hasattr(node, 'name') or node.name is None:
                txt = str(node)
                if txt:
                    para.add_run(txt)
                return
            tag = node.name.lower()
            if tag in ('strong', 'b'):
                for child in node.children:
                    run = para.add_run(child.get_text() if hasattr(child, 'get_text') else str(child))
                    run.bold = True
            elif tag in ('em', 'i'):
                for child in node.children:
                    run = para.add_run(child.get_text() if hasattr(child, 'get_text') else str(child))
                    run.italic = True
            else:
                for child in node.children:
                    add_run_formatted(para, child)

        def process_elem(elem):
            if not hasattr(elem, 'name') or elem.name is None:
                return
            tag = elem.name.lower()
            if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
                level = int(tag[1])
                h = doc.add_heading(level=min(level, 4))
                for child in elem.children:
                    add_run_formatted(h, child)
            elif tag == 'p':
                para = doc.add_paragraph()
                for child in elem.children:
                    add_run_formatted(para, child)
            elif tag in ('ul', 'ol'):
                for li in elem.find_all('li', recursive=False):
                    style_name = 'List Bullet' if tag == 'ul' else 'List Number'
                    try:
                        para = doc.add_paragraph(style=style_name)
                    except Exception:
                        para = doc.add_paragraph()
                    for child in li.children:
                        add_run_formatted(para, child)
            elif tag == 'table':
                rows = elem.find_all('tr')
                if not rows:
                    return
                max_cols = max(len(r.find_all(['td', 'th'])) for r in rows)
                if max_cols == 0:
                    return
                tbl = doc.add_table(rows=len(rows), cols=max_cols)
                tbl.style = 'Table Grid'
                for r_idx, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for c_idx, cell in enumerate(cells):
                        if c_idx < max_cols:
                            wc = tbl.rows[r_idx].cells[c_idx]
                            wc.text = ''
                            para = wc.paragraphs[0]
                            is_header = cell.name == 'th'
                            for child in cell.children:
                                add_run_formatted(para, child)
                            if is_header:
                                for run in para.runs:
                                    run.bold = True

        for elem in soup.children:
            process_elem(elem)

    # — Respuesta —
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    filename = f'{slugify(servicio.titulo) or "servicio"}.docx'
    response = HttpResponse(
        buffer,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


@login_required
def obtener_imagenes_view(request, servicio_id):
    servicio = get_object_or_404(Servicio, id=servicio_id, activo=True)
    if not (servicio.publico or servicio.creado_por == request.user):
        return JsonResponse({'error': 'Sin permisos.'}, status=403)
    imagenes = servicio.imagenes.all()
    return JsonResponse({
        'success': True,
        'imagenes': [{
            'id': img.id,
            'url': img.imagen.url if img.imagen else '',
            'descripcion': img.descripcion or '',
            'fecha_subida': img.fecha_subida.strftime('%d/%m/%Y %H:%M'),
        } for img in imagenes]
    })


@login_required
def subir_imagenes_view(request, servicio_id):
    servicio = get_object_or_404(Servicio, id=servicio_id, activo=True)
    if servicio.creado_por != request.user:
        return JsonResponse({'error': 'Solo puedes agregar imágenes a tus servicios.'}, status=403)

    imagenes_subidas = request.FILES.getlist('imagenes')
    if not imagenes_subidas:
        return JsonResponse({'error': 'No se proporcionaron imágenes.'}, status=400)

    imagenes_creadas = []
    for imagen_file in imagenes_subidas:
        try:
            from PIL import Image
            img = Image.open(imagen_file)
            img.verify()
            imagen_file.seek(0)
        except Exception:
            continue
        obj = ServicioImagen(servicio=servicio, imagen=imagen_file)
        obj.save()
        imagenes_creadas.append({'id': obj.id, 'url': obj.imagen.url})

    if not imagenes_creadas:
        return JsonResponse({'error': 'No se pudieron procesar las imágenes.'}, status=400)

    return JsonResponse({
        'success': True,
        'message': f'{len(imagenes_creadas)} imagen(es) subida(s) correctamente.',
        'imagenes': imagenes_creadas
    })


@login_required
@require_http_methods(["POST"])
def eliminar_imagen_view(request, imagen_id):
    imagen = get_object_or_404(ServicioImagen, id=imagen_id)
    if imagen.servicio.creado_por != request.user:
        return JsonResponse({'error': 'Sin permisos.'}, status=403)
    if imagen.imagen:
        imagen.imagen.delete(save=False)
    imagen.delete()
    return JsonResponse({'success': True, 'message': 'Imagen eliminada correctamente.'})


@login_required
@require_http_methods(["POST"])
def actualizar_descripcion_imagen_view(request, imagen_id):
    imagen = get_object_or_404(ServicioImagen, id=imagen_id)
    if imagen.servicio.creado_por != request.user:
        return JsonResponse({'error': 'Sin permisos.'}, status=403)
    try:
        data = json.loads(request.body)
        imagen.descripcion = data.get('descripcion', '').strip()
        imagen.save(update_fields=['descripcion'])
        return JsonResponse({'success': True})
    except json.JSONDecodeError:
        return JsonResponse({'error': 'JSON inválido.'}, status=400)


@login_required
@require_http_methods(["POST"])
def actualizar_cantidad_view(request, servicio_id):
    servicio = get_object_or_404(Servicio, id=servicio_id, activo=True)
    if servicio.creado_por != request.user:
        return JsonResponse({'success': False, 'error': 'Sin permisos'}, status=403)
    try:
        data = json.loads(request.body)
        cantidad = data.get('cantidad', '').strip()
        if len(cantidad) > 10:
            cantidad = cantidad[:10]
        servicio.cantidad = cantidad if cantidad else None
        servicio.save(update_fields=['cantidad'])
        return JsonResponse({'success': True, 'cantidad': servicio.cantidad or ''})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@login_required
@require_http_methods(["POST"])
def actualizar_actividad_view(request, servicio_id, actividad_idx):
    servicio = get_object_or_404(Servicio, id=servicio_id, activo=True)
    if servicio.creado_por != request.user:
        return JsonResponse({'success': False, 'error': 'Sin permisos'}, status=403)
    try:
        data = json.loads(request.body)
        actividades = list(servicio.actividades_adicionales or [])
        if actividad_idx < 0 or actividad_idx >= len(actividades):
            return JsonResponse({'success': False, 'error': 'Índice inválido'}, status=400)
        act = actividades[actividad_idx]
        if 'cantidad' in data:
            act['cantidad'] = data['cantidad'].strip()
        if 'nombre' in data:
            act['nombre'] = data['nombre'].strip()
        if 'mostrar' in data:
            act['mostrar'] = bool(data['mostrar'])
        servicio.actividades_adicionales = actividades
        servicio.save(update_fields=['actividades_adicionales'])
        return JsonResponse({'success': True, 'actividad': act})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@login_required
@require_http_methods(["POST"])
def actualizar_mostrar_view(request, servicio_id):
    servicio = get_object_or_404(Servicio, id=servicio_id, activo=True)
    if servicio.creado_por != request.user:
        return JsonResponse({'success': False, 'error': 'Sin permisos'}, status=403)
    try:
        data = json.loads(request.body)
        servicio.mostrar = bool(data.get('mostrar', True))
        servicio.save(update_fields=['mostrar'])
        return JsonResponse({'success': True, 'mostrar': servicio.mostrar})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@login_required
@require_http_methods(['POST'])
def extraer_pdf_view(request):
    pdf_file = request.FILES.get('pdf')
    if not pdf_file:
        return JsonResponse({'error': 'No se proporcionó archivo PDF.'}, status=400)
    if pdf_file.size > 5 * 1024 * 1024:
        return JsonResponse({'error': 'El PDF no debe superar 5 MB.'}, status=400)

    texto = ''
    try:
        import pdfplumber
        from io import BytesIO
        with pdfplumber.open(BytesIO(pdf_file.read())) as pdf:
            texto = '\n'.join(page.extract_text() or '' for page in pdf.pages)
    except ImportError:
        try:
            from pypdf import PdfReader
        except ImportError:
            from PyPDF2 import PdfReader
        try:
            reader = PdfReader(pdf_file)
            texto = '\n'.join(page.extract_text() or '' for page in reader.pages)
        except Exception as e:
            return JsonResponse({'error': f'No se pudo leer el PDF: {e}'}, status=500)
    except Exception as e:
        return JsonResponse({'error': f'Error procesando PDF: {e}'}, status=500)

    texto = texto.strip()
    if not texto:
        return JsonResponse({'error': 'El PDF no contiene texto extraíble (puede ser un PDF de imágenes).'}, status=422)

    return JsonResponse({'texto': texto[:8000], 'chars': len(texto)})


@login_required
def obtener_actividades_view(request, servicio_id):
    servicio = get_object_or_404(
        Servicio.objects.select_related('creado_por'),
        id=servicio_id, activo=True
    )
    if not (servicio.publico or servicio.creado_por == request.user):
        return JsonResponse({'error': 'Sin permisos.'}, status=403)
    actividades = servicio.actividades_adicionales or []
    return JsonResponse({
        'success': True,
        'actividades': [{
            'id': i,
            'nombre': a.get('nombre', ''),
            'unidad_medida': a.get('unidad_medida', ''),
            'cantidad': a.get('cantidad', ''),
            'mostrar': a.get('mostrar', False),
        } for i, a in enumerate(actividades)],
        'servicio': {
            'id': servicio.id,
            'titulo': servicio.titulo,
            'unidad_medida': servicio.unidad_medida or '',
            'cantidad': servicio.cantidad or '',
            'mostrar': servicio.mostrar,
            'es_propietario': servicio.creado_por == request.user,
        },
    })
